"""Document parsing and text-normalisation helpers."""
from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Dict


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract readable text from PDF, DOCX, or plain-text files."""
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(file_bytes))
            pages = [(page.extract_text() or "") for page in reader.pages]
            return "\n".join(pages)
        except Exception as exc:  # pragma: no cover - depends on malformed PDFs
            raise ValueError(f"Could not read this PDF: {exc}") from exc
    if suffix == ".docx":
        try:
            from docx import Document

            document = Document(io.BytesIO(file_bytes))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    paragraphs.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(paragraphs)
        except Exception as exc:  # pragma: no cover - depends on malformed DOCX
            raise ValueError(f"Could not read this DOCX: {exc}") from exc
    if suffix in {".txt", ".md"}:
        return file_bytes.decode("utf-8", errors="ignore")
    raise ValueError("Supported formats are PDF, DOCX, TXT and MD.")


def clean_text(text: str) -> str:
    """Normalise whitespace while preserving paragraph boundaries."""
    text = text.replace("\x00", " ").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def word_tokens(text: str) -> list[str]:
    """Return lower-case word and technology tokens for matching."""
    return re.findall(r"[a-zA-Z][a-zA-Z0-9+#.\-/]*", text.lower())


def split_sections(text: str) -> Dict[str, str]:
    """Best-effort section detection for common resume headings."""
    heading_map = {
        "summary": "summary",
        "objective": "summary",
        "profile": "summary",
        "experience": "experience",
        "work experience": "experience",
        "employment": "experience",
        "education": "education",
        "skills": "skills",
        "technical skills": "skills",
        "projects": "projects",
        "certifications": "certifications",
        "achievements": "achievements",
    }
    lines = [line.strip() for line in clean_text(text).splitlines() if line.strip()]
    sections: Dict[str, list[str]] = {"header": []}
    current = "header"
    for line in lines:
        key = re.sub(r"[^a-z ]", "", line.lower()).strip()
        if key in heading_map:
            current = heading_map[key]
            sections.setdefault(current, [])
        else:
            sections.setdefault(current, []).append(line)
    return {key: "\n".join(value).strip() for key, value in sections.items()}


def estimate_contact_fields(text: str) -> dict[str, bool]:
    """Detect useful resume completeness signals without storing personal data."""
    lowered = text.lower()
    return {
        "email": bool(re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text)),
        "phone": bool(re.search(r"(?:\+?\d[\d ()-]{7,}\d)", text)),
        "linkedin": "linkedin.com" in lowered,
        "github": "github.com" in lowered,
    }
